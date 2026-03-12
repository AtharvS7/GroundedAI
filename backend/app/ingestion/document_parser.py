"""Document parser for PDF, DOCX, and TXT files.

Extracts text and metadata page by page. Each format has its own
dedicated parser function. All parsers produce a DocumentObject.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional
from uuid import uuid4

import chardet
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.ingestion.models import DocumentObject, PageObject

logger = logging.getLogger(__name__)


class UnsupportedFormatError(Exception):
    """Raised when the file format is not supported."""
    pass


class ParseError(Exception):
    """Raised when document parsing fails."""
    pass


def parse_document(
    filename: str,
    file_bytes: bytes,
    file_type: Optional[str] = None,
) -> DocumentObject:
    """Parse a document and return a structured DocumentObject.

    Args:
        filename: Original filename.
        file_bytes: Raw file content.
        file_type: File type override (pdf, docx, txt).

    Returns:
        DocumentObject with extracted text and metadata.

    Raises:
        UnsupportedFormatError: If file type is not supported.
        ParseError: If parsing fails.
    """
    if file_type is None:
        ext = Path(filename).suffix.lower().lstrip(".")
        file_type = ext

    if file_type not in ("pdf", "docx", "txt"):
        raise UnsupportedFormatError(
            f"Unsupported file format: {file_type}. "
            "Supported formats: pdf, docx, txt"
        )

    try:
        if file_type == "pdf":
            return _parse_pdf(filename, file_bytes)
        elif file_type == "docx":
            return _parse_docx(filename, file_bytes)
        else:
            return _parse_txt(filename, file_bytes)
    except (UnsupportedFormatError, ParseError):
        raise
    except Exception as e:
        logger.error(f"Failed to parse {filename}: {e}")
        raise ParseError(f"Failed to parse {filename}: {e}") from e


def _parse_pdf(filename: str, file_bytes: bytes) -> DocumentObject:
    """Extract text from PDF using PyMuPDF, page by page."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text").strip()
        pages.append(
            PageObject(
                page_number=page_num + 1,
                raw_text=text,
            )
        )
    doc.close()

    return DocumentObject(
        document_id=uuid4(),
        source_filename=filename,
        file_type="pdf",
        total_pages=len(pages),
        pages=pages,
    )


def _parse_docx(filename: str, file_bytes: bytes) -> DocumentObject:
    """Extract text from DOCX using python-docx.

    Since DOCX doesn't have native pages, we treat each paragraph group
    as page 1 (entire document = 1 logical page). Tables are also extracted.
    """
    doc = DocxDocument(io.BytesIO(file_bytes))

    # Extract paragraphs
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                texts.append(row_text)

    full_text = "\n\n".join(texts)
    pages = [PageObject(page_number=1, raw_text=full_text)]

    return DocumentObject(
        document_id=uuid4(),
        source_filename=filename,
        file_type="docx",
        total_pages=1,
        pages=pages,
    )


def _parse_txt(filename: str, file_bytes: bytes) -> DocumentObject:
    """Extract text from plain text file with encoding detection."""
    # Detect encoding
    detection = chardet.detect(file_bytes)
    encoding = detection.get("encoding", "utf-8") or "utf-8"

    try:
        text = file_bytes.decode(encoding).strip()
    except (UnicodeDecodeError, LookupError):
        text = file_bytes.decode("utf-8", errors="replace").strip()

    pages = [PageObject(page_number=1, raw_text=text)]

    return DocumentObject(
        document_id=uuid4(),
        source_filename=filename,
        file_type="txt",
        total_pages=1,
        pages=pages,
    )
