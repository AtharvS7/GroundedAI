"""Test configuration and fixtures."""

from __future__ import annotations

import os
import sys
import pytest

# Ensure the backend directory is in path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Set test environment variables
os.environ.setdefault("SUPABASE_URL", "https://test.supabase.co")
os.environ.setdefault("SUPABASE_ANON_KEY", "test-anon-key")
os.environ.setdefault("SUPABASE_SERVICE_ROLE_KEY", "test-service-key")
os.environ.setdefault("OLLAMA_BASE_URL", "http://localhost:11434")
os.environ.setdefault("FAISS_INDEX_PATH", "./test_data/faiss_index")
os.environ.setdefault("REDIS_URL", "redis://localhost:6379")
os.environ.setdefault("APP_ENV", "test")


@pytest.fixture
def sample_pdf_bytes():
    """Create a simple test PDF in memory."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "This is a test PDF document.")
    page.insert_text((50, 130), "It contains sample text for testing.")
    page.insert_text((50, 160), "GroundedAI uses this for unit tests.")
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def sample_txt_bytes():
    """Create sample text file bytes."""
    text = (
        "GroundedAI is a production-grade RAG system. "
        "It uses FAISS for vector search and Mistral 7B for generation. "
        "The system eliminates hallucinations by grounding answers in documents. "
        "Each answer includes citations to the source document and page number."
    )
    return text.encode("utf-8")


@pytest.fixture
def sample_docx_bytes():
    """Create a simple test DOCX in memory."""
    from docx import Document
    import io
    doc = Document()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It contains sample text for testing GroundedAI.")
    doc.add_paragraph("The chunker should split this into appropriate pieces.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
